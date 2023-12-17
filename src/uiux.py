"""
"""

# External modules
import PySimpleGUI as sg
import os
import sys
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
# Local modules
from dishdeck_dataclasses import *
from recipebook import RecipeBook
from pantry import Pantry

sg.theme('dark')

RECIPE_FOLDER = 'data/recipes/'
PANTRY_FILE = 'data/pantry.json'
THUMB_FOLDER = 'data/thumb/'
GROCERY_FILE = 'data/grocerylist.docx'

## Set the cwd to the folder above src
script_directory = os.path.dirname(os.path.abspath(__file__))
src_directory = os.path.dirname(script_directory)
os.chdir(src_directory)

# Create master RecipeBook object
recipe_book = RecipeBook()

# Load all recipes saved in local recipes folder
file_names = os.listdir(RECIPE_FOLDER)
for file_name in file_names:#
	# TODO: Make sure there are no duplicate recipes
	recipe_book.import_recipe(RECIPE_FOLDER + file_name)

# Create master Pantry object
pantry = Pantry()

# Import existing pantry ingredients, if there are any
try:
	pantry.load_ingredients(PANTRY_FILE)
except FileNotFoundError:
	print('No pantry file found.')
data_pantry = [[ing.name, ing.quantity if ing.quantity != None else 'N/A', ing.unit if ing.unit != None else 'N/A'] for ing in pantry.get_all_ingredients()]

# Each recipe preview consists of a thumbnail, and as much of the
# description as will fit in the remaining space in the element
def create_recipe_preview(recipe_name: str):
	# TODO: maybe check to make sure the recipe exists in the book
	checkbox_row = sg.Checkbox(
		recipe_name.title(), 
		font=["consolas", 15, "bold"], 
		pad=(0, 10),
		enable_events=True,
		key=recipe_name + "check"
		)

	# Default to missing image thumbnail
	# If a thumbnail is found for the recipe, use that instead
	thumb_name = "Missing Image.png"
	file_names = os.listdir(THUMB_FOLDER)
	for file_name in file_names:
		base_name, ext = os.path.splitext(file_name)
		if base_name == recipe_name:
			thumb_name = file_name
			break

	# Add folder path for full path to thumbnail
	thumb_path = THUMB_FOLDER + thumb_name
	description_thumbnail = sg.Image(filename=(thumb_path), enable_events=True, key=f'_{recipe_name}_RECIPETHUMB_')
	description_text = sg.Text(recipe_book.get_recipe_desc(recipe_name), size=(50, 9), justification="left", enable_events=True, key=f'_{recipe_name}_RECIPETEXT_')
	# Each call creates a new object, avoids duplicate element restriction in PySimpleGUI
	separator = lambda: sg.HorizontalSeparator(color="#FFFFFF", pad=(0, 10)) 
	return sg.pin(sg.Column(
		[[separator()], [checkbox_row], [description_thumbnail, description_text], [separator()]],
		vertical_alignment="top",
		expand_x=True,
		key=recipe_name
		))

# Individual elements of the UI
search_bar = sg.Input(do_not_clear=True, size=(100, 1), pad=((0, 0), (20, 0)), enable_events=True, key="_SEARCH_")
ingredient_list = sg.Listbox([], size=(500, 16), pad=((0, 0), (20, 0)), select_mode="multiple", font=["consolas", 10], no_scrollbar=True, key="_INGREDIENTS_")
recipe_list = sg.Listbox([], size=(500, 16), pad=((0, 0), (20, 0)), select_mode="multiple", font=["consolas", 10], enable_events = True, no_scrollbar=True, key="_SELRECIPELIST_")

# Layout for the recipebook tab
layout_recipebook = [[
	sg.Column(
		[[search_bar], [
			sg.Column(
				[[create_recipe_preview(recipe_name)] for recipe_name in recipe_book.get_recipe_names()],
				scrollable=True,
				vertical_scroll_only=True,
				pad=((0, 10), (10, 0)),
				size=(None, 600),
				expand_x=True,
				key="_RECIPELIST_"
			)
		]],
		pad=((0, 10), (0, 0)),
		vertical_alignment="top",
		size=(700, None),
		expand_y=True
	),
	sg.Column(
		[[ingredient_list], [recipe_list], [sg.Button('Export Groceries', pad=(0, 10), enable_events=True, key="_EXPORTGROCERIES_")]],
		element_justification='c',
		vertical_alignment='t',
		pad=((0, 0), (0, 0)),
	)
]]

# Layout for the pantry tab
layout_pantry = [
		[
			sg.Text('Ingredient', size=(29, 1), pad=((5, 0), (20, 0))),
			sg.Text('Quantity', size=(29, 1), pad=((0, 0), (20, 0))),
			sg.Text('Unit (Cups, oz, etc.)', size=(30, 1), pad=((0, 0), (20, 0)))
		],
		[
			sg.Input(do_not_clear=True, size=(30, 1), pad=((5, 20), (0, 0)), key="_PANTRYADD_ING_"),
   			sg.Input(do_not_clear=True, size=(30, 1), pad=((0, 20), (0, 0)), key="_PANTRYADD_QTY_"),
			sg.Input(do_not_clear=True, size=(30, 1), pad=((0, 20), (0, 0)), key="_PANTRYADD_UNIT_"),
			sg.Button('Add', size=(15, 1), pad=((0, 10), (0, 0)), enable_events=True, key="_PANTRYADD_SUBMIT_"),
			sg.Button('Save Pantry', size=(15, 1), enable_events=True, key="_PANTRY_SAVE_")
		],
		[sg.Table(values=data_pantry, headings=['Ingredient', 'Quantity', 'Unit'], enable_click_events=True, expand_x=True, expand_y=True, key="_PANTRY_")]
]

# Layout for the add recipe tab

layout_addrecipe = [
	[
		sg.Column([
			[
				sg.Image(filename='data/thumb/Missing Image.png', enable_events=True, key="_CHANGETHUMB_")
			],
			[
				sg.Text('Name')
			],
			[
				sg.Input(key="_ADDRECIPE_NAME_")
			],
			[
				sg.Text('Tags (Optional)', pad=((0, 0), (15, 0)))
			],
			[
				sg.Text('Tags must be written in this format: #tag1, #tag2, #etc')
			],
			[
				sg.Input(key="_ADDRECIPE_TAGS_")
			],
			[
				sg.Text('Source (Optional)', pad=((0, 0), (15, 0)))
			],
			[
				sg.Input(key="_ADDRECIPE_SOURCE_")
			],
			[
				sg.Text('Description', pad=((0, 0), (15, 0)))
			],
			[
				sg.Multiline(size=(50, 6), key='_ADDRECIPE_DESCRIPTION_')
			]
		], size=(400, 500), pad=((0, 0), (30, 0)), vertical_alignment='t'),
		sg.Column([
			[
				sg.Text('Ingredients')
			],
			[
				sg.Text('Each ingredient must be on it\'s own line.\nThe name, quantity (optional), and unit (optional) must be separated by a \" | \"\n\
Ex: Tomato Paste | 6 | oz    flour     Potatoes | 7')
			],
			[
				sg.Multiline(size=(None, 6), key='_ADDRECIPE_INGREDIENTS_')
			],
						[
				sg.Text('Instructions', pad=((0, 0), (15, 0)))
			],
			[
				sg.Text('Each instruction should be separated by a \" | \"\nEx: Boil water | cook pasta')
			],
			[
				sg.Multiline(size=(None, 14), key='_ADDRECIPE_INSTRUCTIONS_')
			]
		], pad=((40, 0), (30, 0)), vertical_alignment='t')
	],
	[
		sg.Button('Add Recipe', pad=((420, 0), (30, 0)), enable_events=True, key="_ADDRECIPE_BUTTON_")
	]
]

tab_recipebook = sg.Tab(
	"Recipe Book", 
	layout_recipebook,
	key="_TABRECIPEBOOK_"
	)

tab_pantry = sg.Tab(
	"Pantry",
	layout_pantry,
	key="_TABPANTRY_"
	)

tab_addrecipe = sg.Tab(
	"Add Recipe",
	layout_addrecipe,
	key="_TABADDRECIPE_"
)

# Complete layout for the window
layout = [[sg.TabGroup([
	[tab_recipebook,
	tab_pantry,
	tab_addrecipe]],
	font=["consolas", 20, "bold"],
	border_width=0,
	tab_border_width=2,
	expand_x=True,
	expand_y=True)]
]

# Create the Window
window = sg.Window(
	"Dish Deck",
	layout,
	size=(1000, 700), 
)

# Event loop
while True:
	event, values = window.read()

	if values['_SEARCH_'] != '':
		search = values['_SEARCH_'].lower()
		search_terms = search.split(', ')
		matches = recipe_book.search_recipes(search_terms)
		misses = [name for name in recipe_book.get_recipe_names() if name not in matches]
		for key in matches:
			window[key].Update(visible=True)
		for key in misses:
			window[key].Update(visible=False)
	else:
		for key in recipe_book.get_recipe_names():
			window[key].Update(visible=True)

	# Update ingredients list based on selected recipes
	list_ingredients = []
	list_recipes = []
	for name in recipe_book.get_recipe_names():
		if values[name + "check"]:
			# TODO: Somehow indicate if pantry has partial amount of ingredient
			for ing in recipe_book.get_recipe_ingredients(name):
				pantry_ing = pantry.get_ingredient(ing.name)
				if pantry_ing != None:
					if pantry_ing.quantity == None or (pantry_ing.quantity >= ing.quantity and pantry_ing.unit == ing.unit):
						list_ingredients += [f'[x] {ing}']
					else:
						list_ingredients += [f'[ ] {ing}']
				else:
					list_ingredients += [f'[ ] {ing}']
			
			#Adding recipe name to the 'selected list'
			list_recipes.append(name.title())
			
	if list_ingredients != []:
		max_len = 36 # Max length allowed in listbox boundary
		ing = lambda s: s[:s.find('|')] # Returns ingredient name part of string
		qty = lambda s: s[s.find('|'):] # Returns ingredient quantity part of string
		# Pad quantity parts of ingredient strings to equal length
		list_ingredients = [ing(s) + qty(s).ljust(max(len(qty(s)) for s in list_ingredients)) for s in list_ingredients]
		# Trim any long ingredient strings to listbox width
		list_ingredients = [ing(s)[:(max_len - len(qty(s)) - 4)] + "... " + qty(s) if len(s) > max_len else s for s in list_ingredients]
		# Extend rest of ingredient strings to fill listbox width
		list_ingredients = [ing(s) + (" " * (max_len - len(s))) + qty(s) for s in list_ingredients]

	if event == '_PANTRYADD_SUBMIT_':
		ingredient = values['_PANTRYADD_ING_']
		qty = values['_PANTRYADD_QTY_']
		unit = values['_PANTRYADD_UNIT_']

		if ingredient != '':
			# Accept an entry if both qty and unit are not specified, or just unit is not specified, but not if only unit is specified
			if ingredient != '' and not (qty == '' and unit != ''):
				try:
					pantry.add_ingredient(ingredient, qty, unit)
				except KeyError as error:
					sg.popup_ok(f'Error: {str(error)}')

			# TODO: Display message in UI indicating a submission is invalid

	# Delete an ingredient from the pantry
	if event[0] and event[0] == '_PANTRY_':
		confirm = sg.popup_yes_no(f"Are you sure you want to delete {pantry.get_ingredient(event[2][0])} from the pantry?", title="Confirm deletion")
		if confirm == 'Yes':
			pantry.remove_ingredient(event[2][0])

	if event == '_PANTRY_SAVE_':
		pantry.save_ingredients(PANTRY_FILE)

	#View Recipe
	if "_RECIPETHUMB_" in event or "_RECIPETEXT_" in event:
		split = event.split('_')
		r = split[1] if len(split) > 1 else None		
		recipe = recipe_book.get_recipe(r)
		sg.popup(str(recipe))

	#View recipe card of selected recipe from recipe list
	if event == "_SELRECIPELIST_" and values["_SELRECIPELIST_"] != []:
		sel_recipename = values["_SELRECIPELIST_"][0]
		window['_SEARCH_'].Update(value=sel_recipename)

	if event == "_EXPORTGROCERIES_":
		document = Document()
		document.add_heading('Grocery List', 0)
		all_ingredients = ''
		for ing in list_ingredients:
			all_ingredients += ing + '\n'
		style = document.styles['Normal']
		font = style.font
		font.name = "Consolas"
		document.add_paragraph(all_ingredients, style="Normal")
		document.save(GROCERY_FILE)

	# Add a recipe
	for i in range(1):
		if event == '_ADDRECIPE_BUTTON_':
			name = values['_ADDRECIPE_NAME_']
			tags = values['_ADDRECIPE_TAGS_'].split(', ') if values['_ADDRECIPE_TAGS_'] else None
			source = values['_ADDRECIPE_SOURCE_'] if values['_ADDRECIPE_SOURCE_'] else None
			description = values['_ADDRECIPE_DESCRIPTION_']
			ingredients = values['_ADDRECIPE_INGREDIENTS_']
			instructions = values['_ADDRECIPE_INSTRUCTIONS_']
			if not any([name, description, ingredients, instructions]):
				sg.popup_ok('Error: Required field is empty.')
				break
			if not tags == None and any(tag.startswith('#') == False for tag in tags):
				sg.popup_ok('Error: One or more tags in incorrect format.')
				break
			ingredients = ingredients.split('\n')
			ingredients = [ing.split(' | ') for ing in ingredients]
			for ing in ingredients:
				ing.extend([None] * (3 - len(ing)))
			ingredients = [Ingredient(ing[0], ing[1], ing[2]) for ing in ingredients]
			instructions = instructions.split(' | ')
			recipe_book.add_recipe(RECIPE_FOLDER, name, tags, source, description, ingredients, instructions)
			if sg.popup_ok('Recipe successfully added! Click OK to reload Recipe Book. (or restart the program)') == 'OK':
				python = sys.executable
				os.execv(python, [python, __file__])	

	window["_INGREDIENTS_"].Update(list_ingredients)
	window["_SELRECIPELIST_"].Update(list_recipes)
	window["_PANTRY_"].Update([[ing.name, ing.quantity if ing.quantity != None else 'N/A', ing.unit if ing.unit != None else 'N/A'] for ing in pantry.get_all_ingredients()])
	
	# End program if user closes window
	if event == sg.WIN_CLOSED:
		break

pantry.save_ingredients()
window.close()